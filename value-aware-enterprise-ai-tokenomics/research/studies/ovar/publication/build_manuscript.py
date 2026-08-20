#!/usr/bin/env python3
"""Build the OVAR paper figures and anonymous/identified Word manuscripts.

The document uses a named Springer-LNCS-like override: 15.5 x 23.5 cm page,
12.2 cm text width, Times New Roman 10 pt body, black headings, one column.
The official conference template must still be used for the final submission if
ThinkAI supplies a conference-specific file.
"""

from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REFERENCE_DOCX = ROOT.parent / "action-evidence-safety-research" / "output" / "docx" / "RAER_v2_ThinkAI2026_CAMERA_READY_v1.0.docx"
PUB = ROOT / "publication"
FIG = PUB / "figures"
DOCX = PUB / "output" / "docx"
GATE = ROOT / "calibration" / "results" / "calibration_v1.0" / "calibration_gate.json"

TITLE = "From AI Usage to Auditable Outcomes: A Prospective Negative Calibration of Outcome-Verified AI Resource Allocation"
SHORT = "Prospective Negative Calibration of OVAR"

POLICY_LABELS = {
    "USAGE_ONLY": "Usage only",
    "SELF_REPORTED_VALUE": "Self-report",
    "COST_QUALITY": "Cost-quality",
    "OUTCOME_FLAT": "Outcome-flat",
    "OVAR_LEDGER": "OVAR",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def apply_fixed_geometry(table, widths, total_dxa=6917, indent_dxa=80):
    """Encode matching tblW, tblGrid, and tcW values in exact DXA."""
    raw = [round(w * 1440) for w in widths]
    scale = total_dxa / sum(raw)
    dxas = [round(v * scale) for v in raw[:-1]]
    dxas.append(total_dxa - sum(dxas))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for dxa in dxas:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(dxa))
        grid.append(col)
    for row in table.rows:
        for cell, dxa in zip(row.cells, dxas):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="808080", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after in (
        ("Title", 14, 0, 10),
        ("Heading 1", 12, 10, 4),
        ("Heading 2", 10, 7, 3),
        ("Heading 3", 10, 5, 2),
    ):
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = styles["Caption"]
    cap.font.name = "Times New Roman"
    cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    cap.font.size = Pt(9)
    cap.font.italic = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(5)


def configure_page(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(15.5)
    sec.page_height = Cm(23.5)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)
    sec.top_margin = Cm(1.55)
    sec.bottom_margin = Cm(1.55)
    sec.header_distance = Cm(0.7)
    sec.footer_distance = Cm(0.7)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_text(doc, text, bold_prefix=None, italic=False, center=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.4)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        r = p.add_run(text)
        r.italic = italic
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_equation(doc, equation, number):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = left.add_run(equation)
    run.font.name = "Cambria Math"
    run.font.size = Pt(10)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run(f"({number})")
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tc_pr.append(borders)
    apply_fixed_geometry(table, [4.35, 0.45])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, caption, headers, rows, widths):
    p = doc.add_paragraph(style="Caption")
    p.add_run(caption).bold = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        set_cell_shading(cell, "E7E6E6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(str(h))
        run.bold = True
        run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            para = cells[j].paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(str(value))
            run.font.size = Pt(8)
    set_table_borders(table)
    apply_fixed_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def build_figures(gate):
    FIG.mkdir(parents=True, exist_ok=True)
    summaries = gate["policy_summaries"]
    policies = [s["policy"] for s in summaries]
    labels = [POLICY_LABELS[p] for p in policies]
    colors = ["#7A7A7A", "#A6A6A6", "#5B9BD5", "#70AD47", "#C55A11"]

    metrics = ["false_positive_roi", "false_scale", "false_stop", "authorization_violation"]
    metric_labels = ["False-positive ROI", "False scale", "False stop", "Authorization violation"]
    x = np.arange(len(metrics))
    width = 0.16
    fig, ax = plt.subplots(figsize=(7.2, 4.1))
    for i, s in enumerate(summaries):
        vals = [100 * s["rates"][m] for m in metrics]
        ax.bar(x + (i - 2) * width, vals, width, label=labels[i], color=colors[i])
    ax.set_ylabel("Rate (%)")
    ax.set_xticks(x)
    ax.set_xticklabels(metric_labels, rotation=13, ha="right")
    ax.set_ylim(0, 105)
    ax.grid(axis="y", alpha=0.25)
    ax.legend(ncol=3, fontsize=8, frameon=False, loc="upper center")
    fig.tight_layout()
    fig.savefig(FIG / "policy_error_rates_v1.0.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    weights = [r["measurement_weight"] for r in gate["sensitivity"]]
    fig, ax = plt.subplots(figsize=(7.2, 4.0))
    for i, policy in enumerate(policies):
        losses = [r["losses"][policy] for r in gate["sensitivity"]]
        ax.plot(weights, losses, marker="o", linewidth=1.8, label=labels[i], color=colors[i])
    ax.set_xlabel("Registered measurement-burden weight")
    ax.set_ylabel("Weighted decision loss")
    ax.set_xticks(weights)
    ax.grid(alpha=0.25)
    ax.legend(ncol=2, fontsize=8, frameon=False)
    fig.tight_layout()
    fig.savefig(FIG / "burden_sensitivity_v1.0.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(7.2, 2.5))
    ax.axis("off")
    boxes = [
        (0.02, "Trace + cost\nreconciliation"),
        (0.22, "Outcome contract\n+ baseline"),
        (0.42, "Independent\nevidence"),
        (0.62, "Attribution +\nuncertainty"),
        (0.82, "Decision receipt\nSTOP/REVISE/\nCONTINUE/SCALE"),
    ]
    for x0, label in boxes:
        ax.add_patch(plt.Rectangle((x0, 0.34), 0.16, 0.38, facecolor="#E7E6E6", edgecolor="#404040", linewidth=1))
        ax.text(x0 + 0.08, 0.53, label, ha="center", va="center", fontsize=8)
    for x0, _ in boxes[:-1]:
        ax.annotate("", xy=(x0 + 0.20, 0.53), xytext=(x0 + 0.16, 0.53), arrowprops=dict(arrowstyle="->", lw=1.2))
    ax.text(0.5, 0.12, "Authorization and risk records constrain the final action; token volume is not treated as value.", ha="center", fontsize=8)
    fig.tight_layout()
    fig.savefig(FIG / "ovar_ledger_workflow_v1.0.png", dpi=300, bbox_inches="tight")
    plt.close(fig)


def add_figure(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG / filename), width=Cm(11.8))
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(caption)


def clear_reference_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def manuscript(identified: bool, gate):
    doc = Document(REFERENCE_DOCX)
    clear_reference_body(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "OVAR methods and prospective negative calibration"
    doc.core_properties.keywords = "AI resource allocation; AI ROI; token accounting; outcome evidence; authorization; prospective calibration"
    doc.core_properties.author = "Shaik Khaja Nayab Rasool" if identified else "Anonymous"
    doc.core_properties.last_modified_by = doc.core_properties.author
    doc.core_properties.comments = "Identified camera-ready copy" if identified else "Anonymous double-blind review copy"

    title = doc.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(14)
    if identified:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Shaik Khaja Nayab Rasool")
        r.bold = True
        r.font.size = Pt(10)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("KLU University, Moinabad Road, near the TS Police Academy, Aziznagar, Hyderabad, Telangana 500075, India")
        r.font.size = Pt(9)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run("shaik.khajanayabrasool@gmail.com")
        r.font.size = Pt(9)

    p = add_text(doc, "Abstract. Organizations can observe model calls, tokens, latency, and provider charges, yet those records do not establish that an AI-enabled workflow caused a useful outcome. This paper introduces Outcome-Verified AI Resource Allocation (OVAR), an auditable ledger and decision policy that links resource traces to a predefined outcome contract, independently reviewable evidence, a counterfactual baseline, fully loaded cost, attribution uncertainty, authorization state, and an immutable stop/revise/continue/scale receipt. Rather than presenting a successful performance claim, we report a prospectively specified negative calibration. Five policies were frozen and evaluated once on 48 deliberately constructed cases spanning healthcare, financial services, e-commerce, transportation and logistics, cybersecurity, and customer operations. OVAR reduced false-positive ROI classifications to 2/35 and produced no false-scale decisions, compared with severe proxy errors under usage-only, self-reported-value, and cost-quality policies. However, OVAR generated two false stops among 13 safe reference cases and missed two expired authorizations. It passed five of nine mandatory criteria and was strictly dominated by a simpler outcome-flat policy throughout the registered measurement-burden sensitivity range. Failure tracing showed that lexical parsing did not resolve authorization time and scope reliably. The result rejects the prespecified OVAR v1 calibration hypothesis while supporting a narrower conclusion: outcome evidence can prevent consumption proxies from being mistaken for value, but authorization-sensitive allocation requires structured temporal and scoped records rather than unstructured text heuristics. The study contributes a transparent method, leakage-controlled constructed benchmark, prospective decision gate, and preserved failure mechanism; it does not establish field effectiveness or organizational return on investment.", bold_prefix="Abstract.")
    p.paragraph_format.space_after = Pt(5)
    p = add_text(doc, "Keywords: AI resource allocation · AI return on investment · token accounting · outcome evidence · authorization governance · negative results", bold_prefix="Keywords:")
    p.paragraph_format.space_after = Pt(8)

    add_heading(doc, "1 Introduction", 1)
    for para in [
        "Enterprises increasingly meter generative-AI consumption at the level of requests, models, tokens, tools, and users. This visibility is necessary for financial control, but it is not sufficient for deciding whether an AI project should stop, be revised, continue as a pilot, or scale. A token is a technical and billing unit. Its economic meaning depends on model, modality, cache treatment, workflow position, quality, latency, downstream action, and risk. More importantly, a trace that is fully attributable to a project still does not identify the incremental outcome caused by that project. High use can coexist with rework, weak adoption can coexist with high-value edge cases, and a technically correct system can remain unauthorized for its proposed scope.",
        "Recent work already treats tokens as economic primitives, studies marginal token allocation, quantifies consumption in agentic workflows, predicts agent cost, optimizes model selection, and develops enterprise AI value frameworks [1–12]. Observability platforms and standards further support traces, spend allocation, budgets, evaluation, and cost schemas [13–16]. Accordingly, the problem is not the absence of token accounting or AI ROI frameworks. The narrower unresolved problem studied here is whether an action-ready evidence ledger can prevent usage and self-report proxies from being converted into unsupported return-on-investment and scale decisions, while retaining an auditable authorization safeguard.",
        "We investigate OVAR as a decision-accountability bridge. For each candidate project or episode, OVAR reconciles the AI trace to fully loaded cost, records a predefined outcome contract and measurement window, requires reviewable evidence and a baseline, represents attribution and uncertainty, evaluates authorization and risk facts, and emits a deterministic decision receipt. The method is deliberately policy-oriented: it is not a comparison of vendors, foundation models, agent frameworks, or observability products.",
        "The research was designed so that an apparently favorable comparison could still fail. A 24-case engineering pilot established executability but was not treated as confirmatory because its immutable lock followed execution and the investigator authored both cases and rules. A new 48-case calibration was then constructed across six domains. After a packaging revision removed visible design-stratum and case-order shortcuts, two blinded synthetic reviewers found no blocking construct issue. Reference actions were produced separately, five policy implementations were frozen, 25 tests and file hashes were locked, and the policies were executed once. The held-out stage was conditional on nine mandatory criteria.",
        "The calibration did not authorize progression. OVAR sharply reduced proxy-accounting errors and eliminated false-scale decisions, but it falsely stopped two valid in-scope projects and missed two expired authorization records. Outcome-flat therefore achieved lower composite loss with lower measurement burden at every registered burden weight. We report this result as a prospective negative calibration, not as a failed writing exercise or a post-hoc search for a favorable threshold. The central contribution is evidence that a richer governance layer can lose operational value when its semantics are encoded too coarsely.",
    ]:
        add_text(doc, para, indent=True)
    add_text(doc, "The paper makes four bounded contributions:")
    add_bullets(doc, [
        "an explicit separation between resource attribution, outcome verification, causal attribution, and allocation action;",
        "a reproducible OVAR ledger, comparator information sets, deterministic receipts, and prospectively registered calibration gate;",
        "a leakage-controlled 48-case constructed calibration spanning six consequential enterprise domains; and",
        "a preserved negative result showing that temporal and scoped authorization cannot safely be reduced to a lexical safeguard.",
    ])

    add_heading(doc, "2 Related Work and Novelty Boundary", 1)
    add_heading(doc, "2.1 Token economics, cost visibility, and optimization", 2)
    for para in [
        "Token-economics research has expanded beyond per-token price. Zhu connects token generation, consumption, pricing, allocation, and optimization to workflow production and marginal productivity [1]. Chen et al. synthesize token economics across single-agent, multi-agent, ecosystem, and security levels [2], while Zhu frames agentic systems as marginal token allocators balancing benefit, financial cost, latency, and risk [3]. Salim et al. measure where tokens accumulate in an agentic software-development lifecycle [4], and Bai et al. show that agentic token use is highly variable and not monotonically related to task accuracy [5]. These works motivate resource-aware management but do not make token volume a causal measure of organizational value.",
        "A parallel literature optimizes inference cost or quality. FrugalGPT uses cascades to reduce cost while preserving performance [9]; routing and prompt-compression studies similarly allocate model or context resources under quality constraints [10, 11]. Such work asks how to produce an acceptable model output more efficiently. OVAR instead asks whether an observed project outcome is sufficiently evidenced and attributable to justify a portfolio action. These questions are complementary: a cost-efficient call can support a valueless project, and an expensive call can support a valuable but narrowly used workflow.",
        "Practice standards already cover many claimed components of enterprise tokenomics. FinOps guidance addresses AI allocation, forecasting, unit economics, budgets, and value alignment [12–14]. OpenTelemetry provides generative-AI trace semantics [15], while NIST AI RMF and ISO/IEC 42001 connect organizational objectives, evidence, risk, and governance [16, 17]. OVAR consumes these records; it does not claim to invent observability, cost allocation, governance, routing, or evaluation.",
    ]:
        add_text(doc, para, indent=True)
    add_heading(doc, "2.2 Outcome and causal measurement", 2)
    for para in [
        "Field and controlled studies demonstrate that AI effects can be estimated when exposure, comparison, and outcome measurement are designed explicitly. Noy and Zhang used a preregistered experiment to measure productivity and quality effects in professional writing [18]. Brynjolfsson, Li, and Raymond studied deployment in customer support and found heterogeneous productivity effects [19]. Peng et al. evaluated developer task completion under controlled Copilot access [20]. These studies illustrate a distinction that cost dashboards cannot resolve: observing AI use and an outcome is not equivalent to estimating the counterfactual outcome without AI.",
        "Recent enterprise frameworks already propose expected ROI, closed-loop AI value measurement, stage gates, risk adjustment, and portfolio selection [6–8]. The formal novelty audit therefore rejected broad claims such as the first AI tokenomics framework, first return-on-token metric, first hierarchical allocation mechanism, or first AI ROI method. The surviving contribution under test is the joint use of an episode-level trace-to-cost reconciliation, predefined outcome contract, reviewable evidence, explicit baseline, attribution uncertainty, authorization facts, and a versioned receipt, evaluated through false-ROI and incorrect-action rates.",
    ]:
        add_text(doc, para, indent=True)

    add_table(doc, "Table 1. Position relative to the closest research and practice streams.",
              ["Stream", "Already established", "Boundary retained here"],
              [
                  ["Token economics [1–5]", "Consumption, pricing, marginal allocation, agent cost", "No claim that tokens are new economic units"],
                  ["Cost/quality optimization [9–11]", "Routing, cascades, prompt and inference efficiency", "Decision after outcome evidence, not model selection"],
                  ["Enterprise ROI [6–8]", "Expected value, stage gates, closed-loop value", "Auditable episode evidence and policy-error calibration"],
                  ["FinOps/telemetry [12–15]", "Trace, spend, allocation, schemas, budgets", "Uses telemetry as input; does not replace it"],
                  ["Governance [16,17]", "Risk, objectives, controls, continual improvement", "Tests an authorization-sensitive receipt mechanism"],
                  ["Causal outcome studies [18–20]", "Comparison designs and measured productivity", "Requires a baseline; does not estimate a field effect"],
              ], [1.0, 1.8, 2.0])

    add_heading(doc, "3 OVAR Method", 1)
    add_heading(doc, "3.1 Decision object and ledger", 2)
    add_text(doc, "For project or episode i, policy p observes only its registered information set and selects D_i(p) from {STOP, REVISE, CONTINUE_PILOT, SCALE, INDETERMINATE}. The reference action D_i* is derived separately. OVAR stores five linked records: consumption, work, outcome, value, and allocation. Figure 1 shows the information flow. The final receipt binds input hashes, rule version, reasons, action, and receipt hash so that a decision can be reproduced without exposing a hidden reference label.", indent=True)
    add_figure(doc, "ovar_ledger_workflow_v1.0.png", "Fig. 1. OVAR ledger flow. Authorization and risk constrain the action; token consumption is an input to cost, not a value label.")
    add_table(doc, "Table 2. Minimum OVAR ledger records and their role.",
              ["Record", "Core fields", "Decision role"],
              [
                  ["Consumption", "Provider/model, calls, token classes, tools, latency, charge", "Reconcile technical resource use"],
                  ["Work", "Organization, team, project, workflow, episode, accountable owner", "Bind consumption to an operational unit"],
                  ["Outcome", "Predefined metric, window, threshold, evidence, baseline", "Prevent retrospective success definition"],
                  ["Value", "Incremental benefit, full cost, harm, attribution, uncertainty", "Classify economic evidence"],
                  ["Allocation", "Constraints, authorization, action, reasons, hashes", "Create an auditable portfolio decision"],
              ], [0.85, 2.35, 1.6])

    add_heading(doc, "3.2 Formal quantities", 2)
    add_text(doc, "Fully loaded cost aggregates provider charges with infrastructure, tools, integration, evaluation, human review, governance, and rework. Components remain visible because an apparently favorable provider bill can conceal operational cost.", indent=True)
    add_equation(doc, "C_i = C_provider + C_infra + C_tools + C_integration + C_evaluation + C_human + C_governance + C_rework", 1)
    add_text(doc, "For monetary outcomes, incremental net value subtracts the counterfactual, full cost, and expected harm, while A_i expresses attribution confidence between zero and one.", indent=True)
    add_equation(doc, "N_i = A_i [Y_i(1) - Y_i(0)] - C_i - H_i", 2)
    add_text(doc, "Let [N_i^L, N_i^U] denote the registered uncertainty interval and epsilon the practical-equivalence margin. Evidence insufficiency or unresolved authorization prevents a positive classification regardless of the point estimate.", indent=True)
    add_equation(doc, "R_i = POSITIVE if N_i^L > 0; NEGATIVE if N_i^U < 0; NEUTRAL if [N_i^L,N_i^U] lies within ±ε; INDETERMINATE otherwise", 3)
    add_text(doc, "Policy comparison uses component errors plus normalized measurement burden. The weights were frozen before execution and each component is reported separately so that the composite cannot hide a safety trade-off.", indent=True)
    add_equation(doc, "L(p) = 2 FP_ROI(p) + 4 FS(p) + 2 FSTOP(p) + 8 AUTH(p) + 0.5 M(p)", 4)
    add_text(doc, "The study hypothesis required a useful non-dominated position rather than a single accuracy threshold. With loss L and measurement burden M, a comparator q dominates OVAR when both quantities are no greater and at least one is lower.", indent=True)
    add_equation(doc, "q ≺ OVAR iff L(q) ≤ L(OVAR) and M(q) ≤ M(OVAR), with at least one strict inequality", 5)

    add_heading(doc, "3.3 Comparator policies", 2)
    add_table(doc, "Table 3. Registered policy information sets and normalized measurement burden.",
              ["Policy", "Permitted evidence", "Burden"],
              [
                  ["Usage only", "Utilization and budget position", "0.05"],
                  ["Self-report", "Usage plus owner-reported benefit", "0.10"],
                  ["Cost-quality", "Direct cost and technical acceptance/quality", "0.20"],
                  ["Outcome-flat", "Outcome contract, evidence, baseline, attribution, full cost", "0.65"],
                  ["OVAR", "Outcome-flat plus uncertainty, risk, authorization, receipt constraints", "0.80"],
              ], [1.05, 3.05, 0.7])
    add_text(doc, "The burdens are analytical assumptions used to test the decision trade-off; they are not measured labor hours. Strict field whitelists and recursive forbidden-key checks prevented policies from accessing reference labels. Missing approved-budget fields were mapped prospectively to planned direct charges as a neutral implementation proxy. Authorization was derived only from factual authorization records, and risk used a frozen lexical tier. These implementation choices are part of the method under test, not hidden adjustments.", indent=True)

    add_heading(doc, "4 Prospective Calibration Design", 1)
    add_heading(doc, "4.1 Engineering pilot and separation from calibration", 2)
    add_text(doc, "A 24-case pilot across six domains tested schemas, deterministic policy execution, cost reconciliation, leakage barriers, and decision receipts. Twelve engineering tests passed and the numerical dry run was favorable. It was deliberately excluded from evidentiary claims because the investigator authored cases and rules and the immutable closure followed execution. Pilot cases and labels were not reused in calibration. The reviewer workbook saved after pilot closure is a presentation artifact and does not create human validation.", indent=True)

    add_heading(doc, "4.2 Calibration construction", 2)
    add_text(doc, "The calibration contained 48 new cases: eight in each of healthcare, financial services, e-commerce, transportation and logistics, cybersecurity, and customer operations. Each domain covered high verified value with moderate usage, high usage with weak value, hidden fully loaded cost, weak counterfactual evidence, delayed/shared attribution, a genuine authorization constraint, low adoption with credible value, and a revise-versus-indeterminate boundary. Every case supplied three views: reviewer-visible construct facts, policy inputs, and restricted reference data. Only immutable case identifiers joined the views.", indent=True)
    add_table(doc, "Table 4. Constructed calibration coverage.",
              ["Dimension", "Coverage", "Purpose"],
              [
                  ["Cases", "48", "Prospective calibration only"],
                  ["Domains", "6; 8 cases each", "Cross-domain failure exposure"],
                  ["Design strata", "8 per domain", "Balanced difficult decision conditions"],
                  ["Reference actions", "STOP 17; REVISE 6; CONTINUE 1; SCALE 12; INDETERMINATE 12", "All five actions present"],
                  ["Non-positive ROI denominator", "35", "False-positive ROI"],
                  ["Should-not-scale denominator", "35", "False-scale"],
                  ["Safe reference denominator", "13", "False-stop"],
              ], [1.15, 1.65, 2.0])

    add_heading(doc, "4.3 Construct stress testing and reference adjudication", 2)
    add_text(doc, "The first blinded synthetic construct review detected a visible stratum field and a repeated case-order pattern. A clarity-only packaging revision removed the stratum field, remapped identifiers, and preserved all substantive case facts. Two fresh synthetic rechecks then returned PASS with no blocking leakage or scoreability issue. Exact reviewer agreement varied by dimension: it was high for baseline credibility, cost boundary, attribution, and decision realism, but lower for outcome-contract clarity and evidence auditability; all ratings were within one point. Ambiguity remained intentionally present in partial-evidence cases. Because both reviewers were language-model agents, these results are AI-AI consistency and rubric stress testing, not human inter-rater reliability.", indent=True)
    add_text(doc, "A separate reference adjudication used restricted constructed ground truth and did not access policy code, pilot results, or construct reviews. It reconciled incremental value, eight cost components, expected harm, uncertainty, evidence sufficiency, authorization, ROI state, and action. The resulting action distribution appears in Table 4. Constructed ground truth enables deterministic error analysis but does not simulate the full uncertainty of live organizations.", indent=True)

    add_heading(doc, "4.4 Pre-execution lock and one-time gate", 2)
    add_text(doc, "Before reference labels were joined to outputs, the study froze the candidate cases, reference labels, analysis plan, implementation, test files, dependencies, and hashes. Twenty-five implementation and integrity tests passed, including deterministic receipts, whitelist isolation, forbidden-key rejection, cost reconciliation, no filesystem or reference access, and no case-ID branching. The pre-execution lock SHA-256 was f9648e2305d996ea74e73c7bd4736eabc8744434e70bb906c328a6184c66d2f8. Policies were then executed once. No held-out cases were created or opened.", indent=True)
    add_text(doc, "Progression required all nine criteria: successful pre-execution checks; zero OVAR authorization harm; lower false-positive ROI than usage-only and self-report; false-scale no worse than outcome-flat; false-stop within ten percentage points of the best comparator; indeterminate rate no greater than 30%; no lower-loss/lower-burden comparator; at most one serious OVAR error per domain; and no strict OVAR domination throughout burden weights 0.25–1.00. A failure could not be redefined as success by optimizing a different threshold after seeing the labels.", indent=True)

    add_heading(doc, "5 Results", 1)
    add_heading(doc, "5.1 Policy errors and decisions", 2)
    add_text(doc, "Table 5 reports the complete registered summaries. Consumption and self-report policies classified all 35 non-positive references as positive. Cost-quality also produced 35/35 false-positive ROI classifications. Outcome-flat and OVAR reduced that count to 2/35. OVAR produced no false-scale decisions, but its exact-action agreement was 25/48, below outcome-flat's 32/48. OVAR also stopped two of 13 safe references, whereas every comparator had zero false stops.", indent=True)
    rows = []
    for s in gate["policy_summaries"]:
        r = s["rates"]
        rows.append([
            POLICY_LABELS[s["policy"]],
            f"{100*r['false_positive_roi']:.1f}%",
            f"{100*r['false_scale']:.1f}%",
            f"{100*r['false_stop']:.1f}%",
            f"{100*r['authorization_violation']:.1f}%",
            f"{100*r['exact_action']:.1f}%",
            f"{s['weighted_loss']:.3f}",
        ])
    add_table(doc, "Table 5. Registered calibration performance (48 constructed cases).",
              ["Policy", "FP ROI", "False scale", "False stop", "Auth. viol.", "Exact", "Loss"],
              rows, [0.85, 0.62, 0.68, 0.65, 0.65, 0.62, 0.55])
    add_figure(doc, "policy_error_rates_v1.0.png", "Fig. 2. Policy error rates. Denominators differ by metric: 35 non-positive ROI, 35 should-not-scale, 13 safe references, and 48 authorization cases.")

    add_heading(doc, "5.2 Prospective gate", 2)
    add_text(doc, "OVAR passed five of nine criteria and failed four. It passed the pre-execution integrity criterion, the two proxy-comparison criteria, the indeterminate ceiling, and the domain error ceiling. It failed the zero-authorization-harm rule, the false-stop tolerance, the loss/burden non-dominance rule, and the sensitivity non-dominance rule. The correct decision was therefore STOP for OVAR v1, with no held-out benchmark authorization.", indent=True)
    criteria = gate["prospective_gate"]["criteria"]
    gate_rows = [
        ["Pre-execution tests and hashes", "PASS" if criteria["all_preexecution_tests_and_hashes_passed"] else "FAIL"],
        ["Zero OVAR authorization harm", "PASS" if criteria["zero_ovar_authorization_harm"] else "FAIL"],
        ["Lower FP ROI than usage and self-report", "PASS" if criteria["ovar_fpr_lower_than_usage_and_self"] else "FAIL"],
        ["False scale no worse than outcome-flat", "PASS" if criteria["ovar_false_scale_no_worse_than_outcome_flat"] else "FAIL"],
        ["False stop within 10 percentage points of best", "PASS" if criteria["ovar_false_stop_within_10pp_best"] else "FAIL"],
        ["Indeterminate no greater than 30%", "PASS" if criteria["ovar_indeterminate_at_most_30pct"] else "FAIL"],
        ["Not lower-loss/lower-burden dominated", "PASS" if criteria["ovar_not_loss_burden_dominated"] else "FAIL"],
        ["At most one serious error per domain", "PASS" if criteria["max_one_serious_error_per_domain"] else "FAIL"],
        ["Not dominated throughout sensitivity", "PASS" if criteria["not_dominated_throughout_sensitivity"] else "FAIL"],
    ]
    add_table(doc, "Table 6. Prospective OVAR calibration gate.", ["Mandatory criterion", "Outcome"], gate_rows, [4.0, 0.8])

    add_heading(doc, "5.3 Dominance and burden sensitivity", 2)
    add_text(doc, "At the primary burden weight of 0.5, outcome-flat loss was 1.001 and OVAR loss was 1.155; their assumed burdens were 0.65 and 0.80, respectively. Outcome-flat therefore had both lower loss and lower measurement burden. The same relationship held at every registered burden weight from 0.25 to 1.00 (Fig. 3). This finding matters because OVAR's extra governance layer cannot be justified by the absence of false-scale errors alone when it introduces false stops and leaves the same two authorization violations unresolved.", indent=True)
    add_figure(doc, "burden_sensitivity_v1.0.png", "Fig. 3. Registered sensitivity of weighted decision loss to measurement-burden weight. Outcome-flat dominated OVAR at all four weights.")

    add_heading(doc, "5.4 Binding failure cases", 2)
    add_text(doc, "The four binding cases were not random arithmetic mistakes. They exposed two opposite semantic failures in the same lexical authorization rule. In OC-R032 and OC-R037, approval text contained dates that had expired before the August 2026 decision point. The heuristic detected conditional authorization language but did not compare the valid-until date with the decision timestamp, so it permitted action. In OC-R004 and OC-R011, the record authorized the evaluated scope but also described a different excluded scope. The heuristic treated any absent or out-of-scope phrase as project-wide and stopped valid work. The result was simultaneously under-protective for expired approvals and over-protective for mixed-scope documents.", indent=True)
    add_table(doc, "Table 7. Binding OVAR authorization failures.",
              ["Cases", "Reference", "OVAR error", "Mechanism"],
              [
                  ["OC-R032; OC-R037", "Expired approval; STOP", "Allowed continued action", "No date comparison against decision time"],
                  ["OC-R004; OC-R011", "Current in-scope approval; SCALE", "Stopped valid scope", "Excluded secondary scope generalized to whole project"],
              ], [1.05, 1.25, 1.15, 1.35])
    add_text(doc, "The frozen risk tier used coarse lexical indicators and expected harm as a fixed percentage of cost. It did not directly cause the four binding authorization errors, but it is too crude for a production claim. The calibration therefore identifies a method-design defect rather than evidence that authorization safeguards are unnecessary.", indent=True)

    add_heading(doc, "6 Discussion", 1)
    add_heading(doc, "6.1 What the negative gate establishes", 2)
    for para in [
        "The registered OVAR v1 hypothesis was not supported. This statement is stronger and more precise than saying that the entire research failed. The procedure succeeded in exposing a policy defect before held-out construction or deployment. OVAR's favorable false-positive and false-scale results are descriptive signals, but they cannot override mandatory authorization, false-stop, and dominance criteria. The held-out stage remains absent, so the results are neither confirmatory effectiveness evidence nor an estimate of enterprise ROI.",
        "The calibration also shows why consumption dashboards can be dangerous when treated as decision systems. Usage-only, self-report, and cost-quality policies generated false-positive ROI classifications for every non-positive reference. These constructed results do not estimate their real-world error rates, but they demonstrate internal failure modes: activity is mistaken for causation; claimed time savings are not compared with a baseline; technical quality is treated as business value; and provider cost omits integration, review, governance, and rework.",
        "Outcome-flat performed best under the registered composite. Its advantage should not be overstated as a universal recommendation. Outcome-flat still produced two authorization violations because it did not resolve authorization semantics. It simply avoided OVAR's additional false stops and lower-burden penalty. The result says that the tested lexical safeguard added cost without adding sufficient safety, not that authorization should be removed from enterprise allocation.",
    ]:
        add_text(doc, para, indent=True)

    add_heading(doc, "6.2 Practical interpretation across domains", 2)
    add_text(doc, "The decision problem recurs across industries even though the evidence differs. In healthcare, token volume for a forecasting or documentation assistant does not show improved patient flow or clinically acceptable review burden. In financial services, high usage of a document or fraud assistant does not establish net value when correction cost, regulatory scope, and adverse-action controls are omitted. In e-commerce, a personalization or seller-support workflow can raise clicks while reducing margin or violating consent scope. In transportation, route or maintenance recommendations can appear efficient until exception handling, safety review, and operating-window authorization are included. In cybersecurity, rapid triage can save analyst time yet remain bounded by asset-owner approval and production-system scope. The ledger is intended to make these dependencies explicit before resources are scaled.", indent=True)
    add_text(doc, "For managers, the immediate implication is modest: do not allocate AI budgets solely in proportion to tokens, active users, owner narratives, or technical quality. Require a named outcome, measurement window, comparison design, full-cost boundary, evidence maturity, attribution statement, and accountable decision. However, this study does not justify deploying the OVAR v1 policy. Its authorization record must be redesigned before further prospective evaluation.", indent=True)

    add_heading(doc, "6.3 Design requirements for OVAR v2", 2)
    add_text(doc, "A future version should replace free-text authorization classification with structured records containing subject, resource, permitted action, organizational scope, jurisdiction, valid-from, valid-until, revocation state, required signer, and decision timestamp. Scope containment and temporal validity should be evaluated deterministically before any language-model interpretation. Mixed-scope documents should yield multiple scoped records rather than one global label. The decision receipt should show which authorization record governed which action. Risk estimation should likewise separate consequence severity, exposure probability, and cost rather than apply a lexical percentage of project cost.", indent=True)
    add_text(doc, "Such a redesign must not be tuned and declared successful on the 48 exposed cases. OVAR v2 requires newly constructed design cases, a new analysis plan, boundary tests for dates and nested scope, a new immutable lock, and a new prospective gate. The present calibration cases can be used only as regression tests for known defects. This separation prevents repeated optimization on the same labels from masquerading as confirmation.", indent=True)

    add_heading(doc, "7 Threats to Validity, Ethics, and Research Integrity", 1)
    add_table(doc, "Table 8. Limitations and the inferences they restrict.",
              ["Limitation", "What it prevents", "Mitigation / next evidence"],
              [
                  ["Constructed cases", "No field error-rate or ROI estimate", "Operational multi-site study"],
                  ["48 cases; 13 safe references", "Imprecise false-stop behavior", "Larger preregistered sample"],
                  ["Synthetic construct reviewers", "No human inter-rater validity", "Independent domain experts"],
                  ["Deterministic reference labels", "Understates live adjudication uncertainty", "Prospective human adjudication with audit trail"],
                  ["Assumed measurement burden", "No actual labor/cost estimate", "Time-and-motion measurement"],
                  ["Lexical authorization/risk", "No deployment claim", "Structured records and boundary testing"],
                  ["No held-out set", "No confirmatory generalization", "New sealed benchmark after a passing design gate"],
              ], [1.25, 1.55, 2.0])
    for para in [
        "Internal validity is strengthened by separate views, reference isolation, deterministic receipts, frozen weights, pre-execution hashes, and a one-time run. It is limited by designed rather than naturally sampled cases and by the investigator's role in the overall protocol. The synthetic reviewers reduced avoidable wording and shortcut risk but may share model-family biases. Their agreement is not a substitute for independent human domain review.",
        "Construct validity is limited by analytical measurement-burden values and by the compression of complex organizational outcomes into reference actions. External validity is intentionally low: six domains broaden failure exposure but do not constitute a representative enterprise population. Statistical inference to organizations is therefore inappropriate. The exact counts describe only this calibration set.",
        "The study used no human participants, personal data, patients, customer records, or live organizational decisions. All cases were deliberately constructed. Ethics approval and informed consent were therefore not applicable. Nevertheless, the domains are consequential, so the paper avoids deployment guidance and preserves the failed authorization cases. The author declares no competing interests, and no funds, grants, or other external support were received.",
        "Originality controls include a documented 39-source novelty audit, a source and comparison register, a claim-to-evidence ledger, immutable experiment manifests, quotation avoidance, and a required institutional similarity review before submission. These controls support traceability but cannot prove absence of every conceptual overlap. Claims are narrowed when prior art is close.",
    ]:
        add_text(doc, para, indent=True)

    add_heading(doc, "8 Reproducibility, Data Availability, and AI Disclosure", 1)
    if identified:
        add_text(doc, "Code, schemas, reviewer-visible benchmark records, construct-review outputs, pre-execution locks, decision receipts, and calibration results are intended for public deposit in Zenodo. Permanent repository: Zenodo; DOI/URL: to be assigned upon publication/deposit. Code is intended for release under the MIT License; data and documentation are recommended for CC BY 4.0. Restricted files will be reviewed for safe release and clearly separated from reviewer-visible material. The supplementary manifest records the pilot reviewer workbook SHA-256 as 511e974deccf00dc25e95f759e00cbff0923a021fd0158c6ae5e29105a507b24.", indent=True)
    else:
        add_text(doc, "To preserve double-blind review, the identifying repository URL is withheld. Code, schemas, reviewer-visible benchmark records, construct-review outputs, locks, receipts, and calibration results are prepared for archival release after review. Hashes reported in the method permit verification against the eventual deposit. Restricted files will be separated from reviewer-visible material.", indent=True)
    add_text(doc, "Generative-AI systems were used materially for research assistance, including literature-search support, code and document drafting assistance, and two explicitly labeled synthetic construct-review stress tests. No AI system is listed as an author. The human author defined the research question and claim boundaries, approved the protocol, reviewed the sources and artifacts, interpreted the negative gate, and accepts responsibility for the accuracy, originality, integrity, and final submitted text. Synthetic AI-AI agreement is not represented as human validation. An institution-approved similarity report remains a pre-submission requirement; no similarity percentage is estimated here.", indent=True)
    if identified:
        add_heading(doc, "Author Contributions", 2)
        add_text(doc, "Shaik Khaja Nayab Rasool: Conceptualization; Methodology; Software; Validation; Formal analysis; Investigation; Data curation; Writing - original draft; Writing - review and editing; Visualization; Project administration.")
        add_heading(doc, "Declarations", 2)
        add_text(doc, "Funding. No funds, grants, or other external support were received.")
        add_text(doc, "Competing interests. The author declares no competing interests.")
        add_text(doc, "Ethics. This methods study used deliberately constructed cases and no human participants or personal data. Ethics approval and informed consent were not applicable.")

    add_heading(doc, "9 Conclusion", 1)
    add_text(doc, "OVAR was designed to connect AI resource consumption to evidence that can support an accountable stop, revise, continue, or scale decision. In a prospectively locked 48-case constructed calibration, its outcome-evidence requirements prevented most consumption-proxy errors and eliminated false-scale decisions. The registered method nevertheless failed: it produced two false stops, missed two expired authorizations, and was dominated by a lower-burden outcome-flat policy across all registered sensitivity weights. The OVAR v1 calibration hypothesis was therefore not supported, and no held-out or deployment claim is warranted. The constructive finding is a design boundary: evidence-led AI investment decisions require authorization records with explicit time and scope semantics. A future method must encode those semantics structurally and earn progression on new preregistered cases. Transparent stopping at this gate is itself an operational safeguard against converting an attractive dashboard result into an unsupported claim of enterprise value.", indent=True)

    add_heading(doc, "References", 1)
    refs = [
        "1. Zhu, Q.: AI tokenomics: The economics of tokens, computation, and pricing in foundation models. arXiv:2606.24616 (2026).",
        "2. Chen, Y., et al.: Token economics for LLM agents: A dual-view study from computing and economics. arXiv:2605.09104 (2026).",
        "3. Zhu, S.: Agentic AI systems should be designed as marginal token allocators. arXiv:2605.01214 (2026).",
        "4. Salim, M., Latendresse, J., Khatoonabadi, S.H., Shihab, E.: Tokenomics: Quantifying where tokens are used in agentic software engineering. arXiv:2601.14470 (2026).",
        "5. Bai, L., et al.: How do AI agents spend your money? Analyzing and predicting token consumption in agentic coding tasks. arXiv:2604.22750 (2026).",
        "6. Provost, F., Ipeirotis, P.: AI strategy: How to choose what AI product to implement. arXiv:2607.23733 (2026).",
        "7. Polamarasetty, V.K.: Measuring enterprise AI value in the agentic AI era: A closed-loop framework for adoption, decision intelligence, and ROI optimization. SSRN 6986058 (2026).",
        "8. Krishnan, S., Hepp, A., Gandhi, S.: A multi-layer framework for evaluating the return on investment of AI projects. SSRN 6732598 (2026).",
        "9. Chen, L., Zaharia, M., Zou, J.: FrugalGPT: How to use large language models while reducing cost and improving performance. arXiv:2305.05176 (2023).",
        "10. Ong, I., et al.: RouteLLM: Learning to route LLMs with preference data. arXiv:2406.18665 (2024).",
        "11. Jiang, H., et al.: LLMLingua: Compressing prompts for accelerated inference of large language models. arXiv:2310.05736 (2023).",
        "12. FinOps Foundation: FinOps for AI. https://www.finops.org/framework/technology-categories/ai/ (accessed 12 Aug 2026).",
        "13. FinOps Foundation: How to build a generative AI cost and usage tracker. https://www.finops.org/wg/how-to-build-a-generative-ai-cost-and-usage-tracker/ (accessed 12 Aug 2026).",
        "14. FinOps Foundation: Tokenomics: Managing AI value in SaaS model token costs. https://www.finops.org/wg/token-economics-saas/ (accessed 12 Aug 2026).",
        "15. OpenTelemetry: Generative AI semantic conventions. https://opentelemetry.io/docs/specs/semconv/registry/attributes/gen-ai/ (accessed 12 Aug 2026).",
        "16. National Institute of Standards and Technology: Artificial Intelligence Risk Management Framework (AI RMF 1.0). NIST AI 100-1 (2023).",
        "17. International Organization for Standardization: ISO/IEC 42001:2023, Information technology - Artificial intelligence - Management system (2023).",
        "18. Noy, S., Zhang, W.: Experimental evidence on the productivity effects of generative artificial intelligence. Science 381, 187–192 (2023).",
        "19. Brynjolfsson, E., Li, D., Raymond, L.R.: Generative AI at work. Q. J. Econ. 140, 889–942 (2025).",
        "20. Peng, S., Kalliamvakou, E., Cihon, P., Demirer, M.: The impact of AI on developer productivity: Evidence from GitHub Copilot. arXiv:2302.06590 (2023).",
        "21. Lee, J., Kang, H., Cho, A., Baek, E.: Transferability of token usage rights: A design space analysis of generative AI services. arXiv:2604.26683 (2026).",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.45)
        p.paragraph_format.space_after = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(ref)

    return doc


def main():
    gate = json.loads(GATE.read_text())
    build_figures(gate)
    DOCX.mkdir(parents=True, exist_ok=True)
    for identified, name in (
        (False, "OVAR_ThinkAI2026_Anonymous_Review_Manuscript_v1.0.docx"),
        (True, "OVAR_ThinkAI2026_CAMERA_READY_v1.0.docx"),
    ):
        doc = manuscript(identified, gate)
        doc.save(DOCX / name)
    print(json.dumps({
        "anonymous": str(DOCX / "OVAR_ThinkAI2026_Anonymous_Review_Manuscript_v1.0.docx"),
        "identified": str(DOCX / "OVAR_ThinkAI2026_CAMERA_READY_v1.0.docx"),
        "figures": [str(p) for p in sorted(FIG.glob("*.png"))],
    }, indent=2))


if __name__ == "__main__":
    main()
