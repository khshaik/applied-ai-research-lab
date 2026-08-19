#!/usr/bin/env python3
"""Build the anonymous THINKAI/Springer working DOCX from controlled Markdown."""

from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "papers/thinkai-2026/manuscript/manuscript_working_draft.md"
REFERENCE = Path(
    "/Users/81194246/Library/CloudStorage/OneDrive-Pepsico/Personal/AI/Projects/ThinkAI/"
    "action-evidence-safety-research/git-ready/action-evidence-safety-research/papers/"
    "thinkai-2026/manuscript/initial-submission/RAER_ThinkAI2026_Anonymous_Full_Paper_v1.0.docx"
)
OUTPUT = ROOT / (
    "papers/thinkai-2026/manuscript/initial-submission/"
    "VDCM_ThinkAI2026_Anonymous_Full_Paper_v0.2.docx"
)


def set_font(run, size: float = 10, bold: bool | None = None,
             italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clean_inline(text: str) -> str:
    text = text.replace("**", "").replace("`", "")
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text.strip()


def set_cell_margins(cell, top=35, start=70, bottom=35, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "000000")
        borders.append(node)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.autofit = False
    usable = 4.80
    weights = []
    for col in range(cols):
        max_len = max(len(row[col]) if col < len(row) else 0 for row in rows)
        weights.append(max(8, min(max_len, 42)))
    widths = [usable * weight / sum(weights) for weight in weights]
    for ri, values in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            cell.width = Inches(widths[ci])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(clean_inline(values[ci] if ci < len(values) else ""))
            set_font(run, 7.5, bold=(ri == 0))
    set_repeat_header(table.rows[0])
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, source_md: Path, alt: str, relative: str) -> None:
    image_path = (source_md.parent / relative).resolve()
    if image_path.suffix.lower() == ".svg":
        png = image_path.with_suffix(".png")
        if png.is_file():
            image_path = png
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(image_path), width=Inches(4.65))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run(clean_inline(alt))
    set_font(run, 8, italic=False)


def add_body_paragraph(doc: Document, text: str, *, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(clean_inline(text))
    set_font(run, 10)


def add_reference(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(clean_inline(text)), 8)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(6.10)
    section.page_height = Inches(9.25)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.87)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    for part in (section.header, section.footer):
        for paragraph in part.paragraphs:
            paragraph.text = ""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0
    for name, size, before, after, bold, italic in (
        ("Heading 1", 12, 12, 6, True, False),
        ("Heading 2", 10, 10, 4, True, False),
        ("Heading 3", 10, 8, 2, True, True),
    ):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    props = doc.core_properties
    props.author = "Anonymous"
    props.last_modified_by = "Anonymous"
    props.title = "Beyond Story Points in AI-Assisted Delivery"
    props.subject = "Anonymous THINKAI 2026 submission"
    props.keywords = "AI-assisted software engineering; delivery capacity"
    props.comments = ""


def build() -> None:
    if not REFERENCE.is_file():
        raise FileNotFoundError(REFERENCE)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    clear_body(doc)
    configure_document(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = clean_inline(lines[0].lstrip("# "))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(title), 14, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("Anonymous submission"), 10)

    i = 1
    paragraph: list[str] = []
    table_rows: list[list[str]] = []
    skip_preamble = True
    in_equation = False
    in_references = False
    equation: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            text = " ".join(part.strip() for part in paragraph)
            if text and not text.startswith("**Document status:") and not text.startswith("**Paper route:") and not text.startswith("**Release boundary:"):
                if in_references:
                    add_reference(doc, text)
                else:
                    add_body_paragraph(doc, text)
            paragraph = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            usable = [row for row in table_rows if not all(re.fullmatch(r"\s*:?-+:?\s*", c or "") for c in row)]
            add_table(doc, usable)
            table_rows = []

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if line.startswith("## Abstract"):
            skip_preamble = False
        if skip_preamble:
            i += 1
            continue
        if in_equation:
            if line == r"\]":
                text = " ".join(equation)
                text = text.replace(r"\text{distribution of active human service required at }t_0", "distribution of active human service required at t0")
                text = text.replace(r"\sum_{w,s}", "Σw,s").replace(r"\qquad", "    ").replace(r"\frac{L(r,t)}{C(r,t)}", "L(r,t) / C(r,t)")
                text = text.replace(r"\text", "").replace("{", "").replace("}", "")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                set_font(p.add_run(text), 10, italic=True)
                equation = []
                in_equation = False
            else:
                equation.append(line)
            i += 1
            continue
        if line == r"\[":
            flush_paragraph(); flush_table(); in_equation = True; i += 1; continue
        if not line:
            flush_paragraph(); flush_table(); i += 1; continue
        image = re.fullmatch(r"!\[([^]]+)\]\(([^)]+)\)", line)
        if image:
            flush_paragraph(); flush_table(); add_figure(doc, SOURCE, image.group(1), image.group(2)); i += 1; continue
        if line.startswith("|") and line.endswith("|"):
            flush_paragraph()
            table_rows.append([cell.strip() for cell in line.strip("|").split("|")])
            i += 1; continue
        flush_table()
        if line.startswith("## "):
            flush_paragraph()
            heading = clean_inline(line[3:])
            if heading == "Abstract":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(0)
                set_font(p.add_run("Abstract. "), 10, bold=True)
            elif heading in {"Declarations", "References"}:
                doc.add_paragraph(heading, style="Heading 1")
                if heading == "References":
                    in_references = True
            else:
                doc.add_paragraph(heading, style="Heading 1")
            i += 1; continue
        if line.startswith("### "):
            flush_paragraph(); doc.add_paragraph(clean_inline(line[4:]), style="Heading 2"); i += 1; continue
        if re.match(r"^- ", line):
            flush_paragraph(); add_body_paragraph(doc, line[2:], style="List Bullet"); i += 1; continue
        if in_references and re.match(r"^\d+\. ", line):
            flush_paragraph(); paragraph.append(line); i += 1; continue
        if re.match(r"^\d+\. ", line):
            flush_paragraph(); add_body_paragraph(doc, re.sub(r"^\d+\. ", "", line), style="List Number"); i += 1; continue
        paragraph.append(line)
        i += 1
    flush_paragraph(); flush_table()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    if len(sys.argv) > 1:
        OUTPUT = Path(sys.argv[1]).resolve()
    build()
